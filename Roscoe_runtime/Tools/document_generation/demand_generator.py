#!/usr/bin/env python3
"""
Demand Letter Generator

Converts markdown demand letters to professional DOCX/PDF documents.
Uses python-docx for precise formatting control matching firm letterhead style.

Usage:
    python demand_generator.py input.md output.docx [--pdf] [--letterhead template.docx]
    
Or import and use programmatically:
    from demand_generator import generate_demand_letter
    result = generate_demand_letter(markdown_content, output_path)
"""

import argparse
import re
import subprocess
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Default paths
DEFAULT_LETTERHEAD = Path(__file__).parent.parent.parent / "forms" / "2021 Whaley Letterhead (1).docx"


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Set cell margins in twips (1/20 of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def parse_markdown_table(table_text: str) -> Tuple[List[str], List[List[str]]]:
    """Parse a markdown table into headers and rows."""
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
    
    if len(lines) < 2:
        return [], []
    
    # Parse header
    header_line = lines[0]
    headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
    
    # Skip separator line (line with dashes)
    rows = []
    for line in lines[2:]:
        if line.startswith('|') or '|' in line:
            cells = [cell.strip() for cell in line.split('|') if cell.strip() or line.count('|') > 2]
            # Handle empty cells
            raw_cells = line.split('|')
            cells = [c.strip() for c in raw_cells[1:-1]] if raw_cells[0] == '' else [c.strip() for c in raw_cells]
            if cells:
                rows.append(cells)
    
    return headers, rows


def add_table_to_doc(doc: Document, headers: List[str], rows: List[List[str]], 
                     bold_header: bool = True, first_col_bold: bool = False):
    """Add a formatted table to the document."""
    if not headers and not rows:
        return
    
    num_cols = len(headers) if headers else len(rows[0]) if rows else 0
    if num_cols == 0:
        return
    
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Add headers
    if headers:
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            if bold_header:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                    paragraph.paragraph_format.space_after = Pt(0)
    
    # Add data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = row.cells[col_idx]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    if first_col_bold and col_idx == 0:
                        for run in paragraph.runs:
                            run.bold = True
    
    return table


def parse_markdown_content(markdown: str) -> List[Dict[str, Any]]:
    """
    Parse markdown content into structured elements.
    
    Returns list of dicts with keys: type, content, level (for headers)
    """
    elements = []
    lines = markdown.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            content = line.lstrip('#').strip()
            elements.append({'type': 'header', 'content': content, 'level': level})
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() in ['---', '***', '___']:
            elements.append({'type': 'hr'})
            i += 1
            continue
        
        # Table
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and ('|' in lines[i] or lines[i].strip().startswith('|')):
                table_lines.append(lines[i])
                i += 1
            elements.append({'type': 'table', 'content': '\n'.join(table_lines)})
            continue
        
        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            list_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                item = lines[i].strip()[2:]
                list_items.append(item)
                i += 1
            elements.append({'type': 'list', 'content': list_items})
            continue
        
        # Numbered list
        if re.match(r'^\d+\.\s', line.strip()):
            list_items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                item = re.sub(r'^\d+\.\s*', '', lines[i].strip())
                list_items.append(item)
                i += 1
            elements.append({'type': 'numbered_list', 'content': list_items})
            continue
        
        # Regular paragraph
        if line.strip():
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and '|' not in lines[i]:
                if lines[i].strip().startswith('- ') or lines[i].strip().startswith('* '):
                    break
                if re.match(r'^\d+\.\s', lines[i].strip()):
                    break
                para_lines.append(lines[i])
                i += 1
            elements.append({'type': 'paragraph', 'content': ' '.join(para_lines)})
            continue
        
        i += 1
    
    return elements


def apply_inline_formatting(paragraph, text: str):
    """Apply bold and italic formatting from markdown syntax."""
    # Pattern to match **bold**, *italic*, and regular text
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
    parts = re.findall(pattern, text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def generate_demand_letter(
    markdown_content: str,
    output_path: Path,
    convert_to_pdf: bool = True,
    letterhead_template: Optional[Path] = None
) -> Dict[str, Any]:
    """
    Generate a professional demand letter DOCX/PDF from markdown content.
    
    Args:
        markdown_content: The markdown text of the demand letter
        output_path: Where to save the output DOCX
        convert_to_pdf: Whether to also create a PDF
        letterhead_template: Optional custom letterhead template
    
    Returns:
        Dict with status, paths, and any errors
    """
    result = {
        "status": "success",
        "docx_path": None,
        "pdf_path": None,
        "errors": []
    }
    
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    try:
        doc = Document()
        
        # Configure document styles
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Configure heading styles
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'Times New Roman'
            heading_style.font.bold = True
            if i == 1:
                heading_style.font.size = Pt(14)
            elif i == 2:
                heading_style.font.size = Pt(13)
            else:
                heading_style.font.size = Pt(12)
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Parse and add content
        elements = parse_markdown_content(markdown_content)
        
        for element in elements:
            if element['type'] == 'header':
                level = min(element['level'], 3)
                para = doc.add_paragraph(style=f'Heading {level}')
                apply_inline_formatting(para, element['content'])
                
            elif element['type'] == 'paragraph':
                para = doc.add_paragraph()
                apply_inline_formatting(para, element['content'])
                
            elif element['type'] == 'table':
                headers, rows = parse_markdown_table(element['content'])
                add_table_to_doc(doc, headers, rows)
                doc.add_paragraph()  # Space after table
                
            elif element['type'] == 'list':
                for item in element['content']:
                    para = doc.add_paragraph(style='List Bullet')
                    apply_inline_formatting(para, item)
                    
            elif element['type'] == 'numbered_list':
                for item in element['content']:
                    para = doc.add_paragraph(style='List Number')
                    apply_inline_formatting(para, item)
                    
            elif element['type'] == 'hr':
                # Add a subtle horizontal line
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(12)
        
        # Save DOCX
        doc.save(output_path)
        result["docx_path"] = str(output_path)
        
        # Convert to PDF if requested
        if convert_to_pdf:
            pdf_result = convert_docx_to_pdf(output_path)
            if pdf_result["success"]:
                result["pdf_path"] = pdf_result["path"]
            else:
                result["errors"].append(f"PDF conversion: {pdf_result['error']}")
        
    except Exception as e:
        result["status"] = "error"
        result["errors"].append(str(e))
    
    return result


def convert_docx_to_pdf(docx_path: Path) -> Dict[str, Any]:
    """Convert DOCX to PDF using LibreOffice."""
    result = {"success": False, "path": None, "error": None}
    docx_path = Path(docx_path)
    
    try:
        cmd = [
            'soffice', '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(docx_path.parent),
            str(docx_path)
        ]
        
        process = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        
        expected_pdf = docx_path.with_suffix('.pdf')
        if expected_pdf.exists():
            result["success"] = True
            result["path"] = str(expected_pdf)
        else:
            result["error"] = f"LibreOffice did not produce PDF. stderr: {process.stderr}"
            
    except FileNotFoundError:
        result["error"] = "LibreOffice not found. Install with: brew install --cask libreoffice"
    except subprocess.TimeoutExpired:
        result["error"] = "PDF conversion timed out"
    except Exception as e:
        result["error"] = str(e)
    
    return result


def main():
    parser = argparse.ArgumentParser(description="Generate demand letter from markdown")
    parser.add_argument("input", help="Path to markdown file")
    parser.add_argument("output", help="Output DOCX path")
    parser.add_argument("--pdf", action="store_true", help="Also generate PDF")
    parser.add_argument("--letterhead", help="Custom letterhead template path")
    
    args = parser.parse_args()
    
    # Read markdown
    with open(args.input, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    letterhead = Path(args.letterhead) if args.letterhead else None
    
    result = generate_demand_letter(
        markdown_content=markdown_content,
        output_path=Path(args.output),
        convert_to_pdf=args.pdf,
        letterhead_template=letterhead
    )
    
    import json
    print(json.dumps(result, indent=2))
    
    if result["status"] == "error":
        exit(1)


if __name__ == "__main__":
    main()

