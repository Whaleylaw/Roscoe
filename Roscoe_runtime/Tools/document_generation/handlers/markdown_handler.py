#!/usr/bin/env python3
"""
Markdown Handler for Document Generation

Converts filled Markdown templates to professional DOCX and PDF documents.
Handles YAML frontmatter, tables, images, and exhibit lists.

This handler is used for agent-filled templates like demand letters and complaints,
where the agent has already filled in the content sections.

Usage:
    from handlers.markdown_handler import process_markdown_template
    
    result = process_markdown_template(
        md_path="/path/to/filled_demand.md",
        context={"firm.name": "Whaley Law Firm"},  # Auto-fill fields
        output_pdf=True
    )
"""

import json
import os
import re
import shutil
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import yaml

# Try to import optional dependencies
try:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        BaseDocTemplate, Frame, Image, PageBreak, PageTemplate,
        Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    )
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


# Base paths
CLAUDE_DOCS = Path(os.environ.get("CLAUDE_DOCS", os.environ.get("ROSCOE_ROOT", str(Path(__file__).resolve().parents[2]))))


def load_firm_config() -> Dict[str, Any]:
    """Load firm configuration for letterhead and styling."""
    config_paths = [
        CLAUDE_DOCS / "Tools" / "document_generation" / "firm_config.json",
        CLAUDE_DOCS / "Roscoe_workflows" / "workflows" / "phase_3_demand" / "workflows" / "draft_demand" / "tools" / "firm_config.json",
    ]
    
    for path in config_paths:
        if path.exists():
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
    
    # Default config
    return {
        "firm_name": "Law Firm",
        "attorney_name": "Attorney Name",
        "address": "123 Main St",
        "city_state_zip": "City, ST 12345",
        "phone": "(000) 000-0000",
        "fax": "(000) 000-0000",
    }


def parse_markdown(md_path: str) -> Dict[str, Any]:
    """
    Parse a markdown file into structured data.
    
    Args:
        md_path: Path to markdown file
    
    Returns:
        Dict with:
            - metadata: YAML frontmatter
            - content: Raw content without frontmatter
            - sections: Dict of section_name -> content
            - tables: List of parsed tables
            - images: List of image paths
    """
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    result = {
        'metadata': {},
        'content': '',
        'sections': {},
        'tables': [],
        'images': [],
    }
    
    # Extract YAML frontmatter
    yaml_match = re.search(r'^---\s*\n(.*?)\n---', content, re.DOTALL)
    if yaml_match:
        try:
            result['metadata'] = yaml.safe_load(yaml_match.group(1)) or {}
        except yaml.YAMLError:
            pass
        content = content[yaml_match.end():]
    
    # Remove HTML comments (instructions)
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)
    result['content'] = content.strip()
    
    # Extract sections (## headings)
    section_pattern = r'^##\s+(.+?)$\n(.*?)(?=^##\s|\Z)'
    for match in re.finditer(section_pattern, content, re.MULTILINE | re.DOTALL):
        section_name = match.group(1).strip()
        section_content = match.group(2).strip()
        result['sections'][section_name] = section_content
    
    # Extract images: ![caption](path)
    image_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    for match in re.finditer(image_pattern, content):
        result['images'].append({
            'caption': match.group(1),
            'path': match.group(2)
        })
    
    # Extract tables (simple markdown tables)
    table_pattern = r'(\|.+\|[\r\n]+\|[-:\|\s]+\|[\r\n]+(?:\|.+\|[\r\n]*)+)'
    for match in re.finditer(table_pattern, content):
        table_text = match.group(1)
        rows = []
        for line in table_text.strip().split('\n'):
            if '---' in line:  # Skip separator row
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
        if rows:
            result['tables'].append(rows)
    
    return result


def replace_placeholders(text: str, context: Dict[str, Any]) -> str:
    """Replace {{placeholder}} patterns in text with context values."""
    def replacer(match):
        key = match.group(1)
        value = context.get(key, "")
        
        # Try nested lookup
        if not value and "." in key:
            parts = key.split(".")
            nested = context
            try:
                for part in parts:
                    if isinstance(nested, dict):
                        nested = nested.get(part, "")
                    else:
                        break
                if nested and not isinstance(nested, dict):
                    value = str(nested)
            except Exception:
                pass
        
        return str(value) if value else f"[{key}]"
    
    return re.sub(r'\{\{([^}]+)\}\}', replacer, text)


def process_markdown_template(
    md_path: str,
    context: Dict[str, Any] = None,
    output_pdf: bool = True,
    output_docx: bool = True
) -> Dict[str, Any]:
    """
    Process a Markdown template and convert to DOCX/PDF.
    
    Args:
        md_path: Path to the markdown file
        context: Additional context for placeholder replacement
        output_pdf: Whether to generate PDF
        output_docx: Whether to generate DOCX
    
    Returns:
        Dict with status, paths, and any errors
    """
    result = {
        "status": "error",
        "docx_path": None,
        "pdf_path": None,
        "fields_filled": [],
        "fields_missing": [],
        "errors": [],
    }
    
    path = Path(md_path)
    
    if not path.exists():
        result["errors"].append(f"File not found: {md_path}")
        return result
    
    context = context or {}
    
    # Load firm config and merge into context
    firm_config = load_firm_config()
    context.setdefault("firm", firm_config)
    context.setdefault("TODAY", datetime.now().strftime("%B %d, %Y"))
    
    try:
        # Parse the markdown
        parsed = parse_markdown(md_path)
        
        # Merge metadata into context
        for key, value in parsed.get('metadata', {}).items():
            if key not in context:
                context[key] = value
        
        # Replace placeholders in content
        content = replace_placeholders(parsed['content'], context)
        
        # Generate output files
        base_path = path.with_suffix('')
        
        if output_docx:
            if PYTHON_DOCX_AVAILABLE:
                docx_path = base_path.with_suffix('.docx')
                create_docx_from_markdown(content, docx_path, context, firm_config)
                result["docx_path"] = str(docx_path)
            else:
                result["errors"].append("python-docx not installed. Run: pip install python-docx")
        
        if output_pdf:
            if REPORTLAB_AVAILABLE:
                pdf_path = base_path.with_suffix('.pdf')
                create_pdf_from_markdown(content, pdf_path, context, firm_config, parsed)
                result["pdf_path"] = str(pdf_path)
            elif result.get("docx_path"):
                # Try to convert DOCX to PDF using LibreOffice
                pdf_path = base_path.with_suffix('.pdf')
                success, error = convert_docx_to_pdf(Path(result["docx_path"]), pdf_path)
                if success:
                    result["pdf_path"] = str(pdf_path)
                else:
                    result["errors"].append(f"PDF conversion failed: {error}")
            else:
                result["errors"].append("reportlab not installed. Run: pip install reportlab")
        
        result["status"] = "success"
        
    except Exception as e:
        result["errors"].append(str(e))
        import traceback
        result["errors"].append(traceback.format_exc())
    
    return result


def create_docx_from_markdown(
    content: str,
    output_path: Path,
    context: Dict[str, Any],
    firm_config: Dict[str, Any]
) -> None:
    """Create a DOCX file from markdown content."""
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx not installed")
    
    doc = Document()
    
    # Add letterhead
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = firm_config.get("firm_name", "")
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process content by lines
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
        else:
            doc.add_paragraph(line)
    
    doc.save(output_path)


def create_pdf_from_markdown(
    content: str,
    output_path: Path,
    context: Dict[str, Any],
    firm_config: Dict[str, Any],
    parsed: Dict[str, Any]
) -> None:
    """Create a PDF file from markdown content using ReportLab."""
    if not REPORTLAB_AVAILABLE:
        raise ImportError("reportlab not installed")
    
    # Create document
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=14,
        spaceAfter=12,
        alignment=TA_CENTER,
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
        alignment=TA_JUSTIFY,
    )
    
    # Build content
    story = []
    
    # Letterhead
    letterhead = f"""
    <b>{firm_config.get('firm_name', '')}</b><br/>
    {firm_config.get('attorney_name', '')}<br/>
    {firm_config.get('address', '')}<br/>
    {firm_config.get('city_state_zip', '')}<br/>
    Phone: {firm_config.get('phone', '')} | Fax: {firm_config.get('fax', '')}
    """
    story.append(Paragraph(letterhead, ParagraphStyle('Letterhead', alignment=TA_CENTER)))
    story.append(Spacer(1, 0.5 * inch))
    
    # Process content
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            story.append(Spacer(1, 6))
        elif line.startswith('# '):
            story.append(Paragraph(line[2:], title_style))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], styles['Heading2']))
        elif line.startswith('### '):
            story.append(Paragraph(line[4:], styles['Heading3']))
        else:
            # Escape special characters for ReportLab
            safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(safe_line, body_style))
    
    doc.build(story)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> Tuple[bool, str]:
    """Convert DOCX to PDF using LibreOffice."""
    try:
        cmd = [
            "soffice",
            "--headless",
            "--convert-to", "pdf",
            str(docx_path),
            "--outdir", str(pdf_path.parent)
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        expected_pdf = pdf_path.parent / (docx_path.stem + ".pdf")
        if expected_pdf.exists() and expected_pdf != pdf_path:
            shutil.move(expected_pdf, pdf_path)
        
        if pdf_path.exists():
            return True, ""
        else:
            return False, f"PDF not created: {result.stderr}"
    
    except FileNotFoundError:
        return False, "LibreOffice not found"
    except subprocess.TimeoutExpired:
        return False, "Conversion timed out"
    except Exception as e:
        return False, str(e)


# =============================================================================
# CLI for testing
# =============================================================================

def main():
    """Test markdown processing."""
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python markdown_handler.py <markdown_path>")
        sys.exit(1)
    
    result = process_markdown_template(sys.argv[1])
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()

