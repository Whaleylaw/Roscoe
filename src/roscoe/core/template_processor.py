"""
Template processor for document generation.

Handles:
- Loading template YAML metadata
- Resolving field values from graph, config, or computed sources
- Merging fields into DOCX templates
- Converting to PDF via LibreOffice
"""

import json
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import yaml
from docx import Document

from roscoe.core.workspace_resolver import GCS_WORKSPACE


TEMPLATES_DIR = "/Templates"


def list_templates() -> List[Dict]:
    """List all available templates with their metadata."""
    templates = []
    templates_path = GCS_WORKSPACE / "Templates"

    if not templates_path.exists():
        return []

    for yaml_file in templates_path.glob("*.yaml"):
        try:
            with open(yaml_file, 'r') as f:
                metadata = yaml.safe_load(f)
                # Check that corresponding .docx exists
                docx_path = yaml_file.with_suffix('.docx')
                if docx_path.exists():
                    metadata['_yaml_path'] = str(yaml_file)
                    metadata['_docx_path'] = str(docx_path)
                    templates.append(metadata)
        except Exception as e:
            print(f"Warning: Could not load template {yaml_file}: {e}")

    return templates


def get_template(template_id: str) -> Optional[Dict]:
    """Get a specific template by ID."""
    for template in list_templates():
        if template.get('id') == template_id:
            return template
    return None


async def resolve_graph_field(query: str, params: Dict) -> Optional[str]:
    """Execute a Cypher query to resolve a field value."""
    from roscoe.core.graphiti_client import run_cypher_query

    try:
        result = await run_cypher_query(query, params)
        if result and len(result) > 0:
            # Return first column of first row
            first_row = result[0]
            if isinstance(first_row, dict):
                # Get the first non-None value
                for val in first_row.values():
                    if val is not None:
                        return str(val)
            return str(first_row) if first_row else None
    except Exception as e:
        print(f"Graph query failed: {e}")
    return None


def get_config_value(field_name: str) -> Optional[str]:
    """Get a value from firm_settings.json."""
    config_path = GCS_WORKSPACE / "Database" / "firm_settings.json"
    try:
        with open(config_path, 'r') as f:
            config = json.load(f)

        # Map field names to config keys
        mapping = {
            'firm_name': 'firm_name',
            'firm_address': lambda c: f"{c.get('address_line1', '')}, {c.get('city', '')}, {c.get('state', '')} {c.get('zip', '')}",
            'firm_phone': 'phone',
            'firm_email': 'email',
            'firm_fax': 'fax',
        }

        if field_name in mapping:
            if callable(mapping[field_name]):
                return mapping[field_name](config)
            return config.get(mapping[field_name])
    except Exception as e:
        print(f"Config lookup failed: {e}")
    return None


def get_computed_value(field_name: str) -> Optional[str]:
    """Get a computed/auto-generated value."""
    if field_name == 'today_date':
        return datetime.now().strftime("%B %d, %Y")
    if field_name == 'today_date_short':
        return datetime.now().strftime("%m/%d/%Y")
    if field_name == 'today_date_iso':
        return datetime.now().strftime("%Y-%m-%d")
    return None


async def resolve_fields(
    template: Dict,
    graph_inputs: Dict[str, str],
    direct_inputs: Dict[str, str]
) -> Dict[str, str]:
    """Resolve all field values for a template."""
    resolved = {}
    input_mode = template.get('input_mode', 'direct')

    for field in template.get('fields', []):
        field_name = field['name']
        source = field.get('source', 'graph' if 'graph_query' in field else 'input')

        if source == 'computed':
            resolved[field_name] = get_computed_value(field_name) or ''
        elif source == 'config':
            resolved[field_name] = get_config_value(field_name) or ''
        elif source == 'input':
            # Direct from graph_inputs or direct_inputs
            resolved[field_name] = graph_inputs.get(field_name) or direct_inputs.get(field_name) or ''
        elif 'graph_query' in field and input_mode == 'graph':
            # Build params from graph_inputs
            params = {**graph_inputs}
            value = await resolve_graph_field(field['graph_query'], params)
            resolved[field_name] = value or direct_inputs.get(field_name, '')
        else:
            # Direct input mode
            resolved[field_name] = direct_inputs.get(field_name, '')

    return resolved


def _replace_in_paragraph(para, fields: Dict[str, str]) -> None:
    """Replace placeholders in a paragraph while preserving formatting."""
    for field_name, value in fields.items():
        placeholder = f"{{{{{field_name}}}}}"
        if placeholder in para.text:
            # Check if placeholder spans multiple runs
            full_text = para.text
            if placeholder in full_text:
                # Try to replace in individual runs first
                replaced = False
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value or '')
                        replaced = True

                # If placeholder spans runs, we need to handle it differently
                if not replaced and placeholder in full_text:
                    # Reconstruct the paragraph text
                    new_text = full_text.replace(placeholder, value or '')
                    # Clear all runs and set text on first run
                    if para.runs:
                        # Keep formatting of first run
                        first_run = para.runs[0]
                        for run in para.runs[1:]:
                            run.text = ''
                        first_run.text = new_text


def merge_docx(template_path: str, fields: Dict[str, str], output_path: str) -> str:
    """Merge fields into DOCX template and save."""
    doc = Document(template_path)

    # Replace in paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, fields)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, fields)

    # Replace in headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for para in header.paragraphs:
                    _replace_in_paragraph(para, fields)
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    _replace_in_paragraph(para, fields)

    doc.save(output_path)
    return output_path


def convert_to_pdf(docx_path: str, output_dir: str) -> str:
    """Convert DOCX to PDF using LibreOffice headless."""
    try:
        # LibreOffice headless conversion
        result = subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ], capture_output=True, text=True, timeout=60)

        if result.returncode != 0:
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")

        # Return path to generated PDF
        pdf_name = Path(docx_path).stem + '.pdf'
        return str(Path(output_dir) / pdf_name)
    except FileNotFoundError:
        raise Exception("LibreOffice not installed. Install with: apt-get install libreoffice")


async def generate_document(
    template_id: str,
    case_name: Optional[str] = None,
    graph_inputs: Optional[Dict[str, str]] = None,
    direct_inputs: Optional[Dict[str, str]] = None,
    output_path: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Generate a document from a template.

    Returns dict with:
    - success: bool
    - pdf_path: workspace-relative path to generated PDF
    - docx_path: workspace-relative path to intermediate DOCX
    - fields_used: dict of resolved field values
    - error: error message if failed
    """
    template = get_template(template_id)
    if not template:
        return {"success": False, "error": f"Template '{template_id}' not found"}

    graph_inputs = graph_inputs or {}
    direct_inputs = direct_inputs or {}

    # Add case_name to graph_inputs if provided
    if case_name:
        graph_inputs['case_name'] = case_name

    # Resolve all fields
    try:
        fields = await resolve_fields(template, graph_inputs, direct_inputs)
    except Exception as e:
        return {"success": False, "error": f"Failed to resolve fields: {e}"}

    # Determine output path
    if not output_path and case_name:
        # Default: save to case correspondence folder
        output_filename = template.get('output_filename', f"{template_id}_{datetime.now().strftime('%Y%m%d')}.pdf")
        # Replace placeholders in filename
        for key, val in fields.items():
            # Sanitize value for filename
            safe_val = str(val or '').replace(' ', '_').replace('/', '-')[:20]
            output_filename = output_filename.replace(f"{{{key}}}", safe_val)
        output_filename = output_filename.replace("{date}", datetime.now().strftime('%Y%m%d'))
        output_path = f"/projects/{case_name}/correspondence/{output_filename}"
    elif not output_path:
        output_filename = f"{template_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        output_path = f"/Reports/{output_filename}"

    # Ensure output directory exists
    abs_output_dir = GCS_WORKSPACE / Path(output_path).parent.relative_to('/')
    abs_output_dir.mkdir(parents=True, exist_ok=True)

    # Create temp directory for processing
    with tempfile.TemporaryDirectory() as temp_dir:
        # Merge fields into DOCX
        temp_docx = str(Path(temp_dir) / "merged.docx")
        try:
            merge_docx(template['_docx_path'], fields, temp_docx)
        except Exception as e:
            return {"success": False, "error": f"Failed to merge template: {e}"}

        # Convert to PDF
        try:
            temp_pdf = convert_to_pdf(temp_docx, temp_dir)
        except Exception as e:
            return {"success": False, "error": f"Failed to convert to PDF: {e}"}

        # Copy PDF to output location
        abs_pdf_path = GCS_WORKSPACE / Path(output_path).relative_to('/')
        shutil.copy2(temp_pdf, abs_pdf_path)

        # Save DOCX too (useful for editing)
        docx_output_path = output_path.replace('.pdf', '.docx')
        abs_docx_path = GCS_WORKSPACE / Path(docx_output_path).relative_to('/')
        shutil.copy2(temp_docx, abs_docx_path)

    return {
        "success": True,
        "pdf_path": output_path,
        "docx_path": docx_output_path,
        "fields_used": fields,
        "template_name": template.get('name'),
    }
