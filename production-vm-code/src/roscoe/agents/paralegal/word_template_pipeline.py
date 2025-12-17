"""
Word Template Pipeline

This module provides utilities for filling Word (.docx) templates with Jinja2-style
placeholders and converting documents to PDF using LibreOffice headless mode.

Key Features:
- Fill .docx templates using docxtpl (Jinja2 syntax: {{ variable_name }})
- Convert legacy .doc files to .docx format
- Export filled templates to high-fidelity PDF
- Context sanitization for various data types
- Cross-platform support (macOS, Linux/Ubuntu)

Dependencies:
- docxtpl: Jinja2 template filling for Word documents
- python-docx: Low-level Word document manipulation
- LibreOffice: External dependency for PDF conversion (must be installed)

Usage:
    from word_template_pipeline import fill_and_export_template
    
    result = fill_and_export_template(
        template_path="/path/to/template.docx",
        output_path="/path/to/output.docx",
        context={"client_name": "John Doe", "date": "2024-01-15"},
        export_pdf=True
    )
"""

import os
import re
import subprocess
import shutil
import tempfile
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from datetime import datetime, date
from decimal import Decimal


# =============================================================================
# CUSTOM EXCEPTIONS
# =============================================================================

class TemplateError(Exception):
    """Base exception for template-related errors."""
    pass


class LibreOfficeNotFoundError(TemplateError):
    """Raised when LibreOffice is not found on the system."""
    pass


class ConversionError(TemplateError):
    """Raised when document conversion fails."""
    pass


class TemplateNotFoundError(TemplateError):
    """Raised when the specified template file is not found."""
    pass


class PlaceholderError(TemplateError):
    """Raised when placeholder resolution fails."""
    pass


# =============================================================================
# LIBREOFFICE DETECTION
# =============================================================================

def _find_libreoffice() -> Optional[str]:
    """
    Find the LibreOffice executable (soffice) across different platforms.
    
    Returns:
        Path to soffice executable, or None if not found.
    """
    # Common paths for LibreOffice
    possible_paths = [
        # macOS paths
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/opt/homebrew/bin/soffice",
        "/usr/local/bin/soffice",
        # Linux paths
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/usr/lib/libreoffice/program/soffice",
        "/snap/bin/libreoffice",
        # Check PATH
        "soffice",
        "libreoffice",
    ]
    
    # First check explicit paths
    for path in possible_paths[:6]:  # Absolute paths
        if os.path.isfile(path) and os.access(path, os.X_OK):
            return path
    
    # Check PATH using shutil.which
    for cmd in ["soffice", "libreoffice"]:
        result = shutil.which(cmd)
        if result:
            return result
    
    return None


def check_dependencies() -> Dict[str, Any]:
    """
    Check if all required dependencies are available.
    
    Returns:
        Dictionary with dependency status:
        {
            "docxtpl_available": bool,
            "python_docx_available": bool,
            "libreoffice_available": bool,
            "libreoffice_path": str or None,
            "all_available": bool,
            "missing": list of missing dependencies
        }
    """
    result = {
        "docxtpl_available": False,
        "python_docx_available": False,
        "libreoffice_available": False,
        "libreoffice_path": None,
        "all_available": False,
        "missing": []
    }
    
    # Check docxtpl
    try:
        import docxtpl
        result["docxtpl_available"] = True
    except ImportError:
        result["missing"].append("docxtpl")
    
    # Check python-docx
    try:
        import docx
        result["python_docx_available"] = True
    except ImportError:
        result["missing"].append("python-docx")
    
    # Check LibreOffice
    lo_path = _find_libreoffice()
    if lo_path:
        result["libreoffice_available"] = True
        result["libreoffice_path"] = lo_path
    else:
        result["missing"].append("LibreOffice (soffice)")
    
    result["all_available"] = len(result["missing"]) == 0
    
    return result


# =============================================================================
# CONTEXT SANITIZATION
# =============================================================================

def sanitize_context(context: Dict[str, Any]) -> Dict[str, Any]:
    """
    Recursively sanitize context values for Jinja2 template rendering.
    
    Handles:
    - None values -> empty string
    - datetime/date objects -> formatted string
    - Decimal objects -> string representation
    - Nested dicts and lists
    
    Args:
        context: Dictionary of template variables
        
    Returns:
        Sanitized dictionary safe for Jinja2 rendering
    """
    def _sanitize_value(value: Any) -> Any:
        if value is None:
            return ""
        elif isinstance(value, datetime):
            return value.strftime("%B %d, %Y")  # e.g., "January 15, 2024"
        elif isinstance(value, date):
            return value.strftime("%B %d, %Y")
        elif isinstance(value, Decimal):
            return str(value)
        elif isinstance(value, dict):
            return {k: _sanitize_value(v) for k, v in value.items()}
        elif isinstance(value, (list, tuple)):
            return [_sanitize_value(item) for item in value]
        else:
            return value
    
    return _sanitize_value(context)


# =============================================================================
# TEMPLATE PLACEHOLDER EXTRACTION
# =============================================================================

def list_template_placeholders(template_path: Union[str, Path]) -> List[str]:
    """
    Extract all Jinja2 placeholders from a .docx template.
    
    Args:
        template_path: Path to the .docx template file
        
    Returns:
        List of unique placeholder names found in the template
        
    Raises:
        TemplateNotFoundError: If template file doesn't exist
        TemplateError: If template cannot be read
    """
    template_path = Path(template_path)
    
    if not template_path.exists():
        raise TemplateNotFoundError(f"Template not found: {template_path}")
    
    try:
        from docxtpl import DocxTemplate
        from docx import Document
        
        # Use python-docx directly for inspection (more reliable than docxtpl.docx)
        doc = Document(str(template_path))
        
        placeholders = set()
        
        # Jinja2 variable patterns: {{ var }}, {{ var.attr }}, {{ var|filter }}
        pattern = r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*(?:\.[a-zA-Z_][a-zA-Z0-9_]*)*)\s*(?:\|[^}]*)?\}\}'
        
        # Check main document paragraphs
        if doc.paragraphs:
            for para in doc.paragraphs:
                if para.text:
                    matches = re.findall(pattern, para.text)
                    placeholders.update(matches)
        
        # Check tables
        if doc.tables:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            matches = re.findall(pattern, cell.text)
                            placeholders.update(matches)
        
        # Check headers and footers
        if doc.sections:
            for section in doc.sections:
                # Headers
                try:
                    if section.header and section.header.paragraphs:
                        for para in section.header.paragraphs:
                            if para.text:
                                matches = re.findall(pattern, para.text)
                                placeholders.update(matches)
                except Exception:
                    pass
                
                # Footers
                try:
                    if section.footer and section.footer.paragraphs:
                        for para in section.footer.paragraphs:
                            if para.text:
                                matches = re.findall(pattern, para.text)
                                placeholders.update(matches)
                except Exception:
                    pass
        
        return sorted(list(placeholders))
        
    except ImportError:
        raise TemplateError("docxtpl is not installed. Run: pip install docxtpl")
    except Exception as e:
        raise TemplateError(f"Failed to read template: {str(e)}")


# =============================================================================
# TEMPLATE RENDERING
# =============================================================================

def render_docx_template(
    template_path: Union[str, Path],
    context: Dict[str, Any],
    output_path: Union[str, Path],
    autoescape: bool = True,
    sanitize: bool = True,
) -> Dict[str, Any]:
    """
    Render a .docx template with the provided context.
    
    Uses docxtpl (Jinja2) for template filling. Placeholders should use
    Jinja2 syntax: {{ variable_name }}
    
    Args:
        template_path: Path to the .docx template
        context: Dictionary of variables to fill in
        output_path: Path where the filled document will be saved
        autoescape: If True, escape XML special characters in values
        sanitize: If True, sanitize context values (None->empty string, etc.)
        
    Returns:
        Dictionary with operation result:
        {
            "success": bool,
            "output_path": str,
            "placeholders_found": list,
            "placeholders_filled": list,
            "error": str (if failed)
        }
        
    Raises:
        TemplateNotFoundError: If template doesn't exist
        TemplateError: If rendering fails
    """
    template_path = Path(template_path)
    output_path = Path(output_path)
    
    if not template_path.exists():
        raise TemplateNotFoundError(f"Template not found: {template_path}")
    
    try:
        from docxtpl import DocxTemplate
        
        # Load template
        doc = DocxTemplate(str(template_path))
        
        # Sanitize context if requested
        if sanitize:
            context = sanitize_context(context)
        
        # Find placeholders before rendering
        placeholders_found = list_template_placeholders(template_path)
        placeholders_filled = [p for p in placeholders_found if p in context]
        
        # Render the template
        doc.render(context, autoescape=autoescape)
        
        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Save the document
        doc.save(str(output_path))
        
        return {
            "success": True,
            "output_path": str(output_path),
            "placeholders_found": placeholders_found,
            "placeholders_filled": placeholders_filled,
            "placeholders_unfilled": [p for p in placeholders_found if p not in context],
        }
        
    except ImportError:
        raise TemplateError("docxtpl is not installed. Run: pip install docxtpl")
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "output_path": None,
        }


# =============================================================================
# DOCUMENT CONVERSION
# =============================================================================

def convert_doc_to_docx(
    input_path: Union[str, Path],
    output_dir: Optional[Union[str, Path]] = None,
    timeout: int = 120,
) -> Dict[str, Any]:
    """
    Convert a legacy .doc file to .docx format using LibreOffice.
    
    Args:
        input_path: Path to the .doc file
        output_dir: Directory for output file (defaults to same as input)
        timeout: Maximum seconds to wait for conversion
        
    Returns:
        Dictionary with operation result:
        {
            "success": bool,
            "output_path": str,
            "error": str (if failed)
        }
        
    Raises:
        LibreOfficeNotFoundError: If LibreOffice is not installed
        ConversionError: If conversion fails
    """
    input_path = Path(input_path)
    
    if not input_path.exists():
        return {
            "success": False,
            "error": f"Input file not found: {input_path}",
            "output_path": None,
        }
    
    # Find LibreOffice
    soffice_path = _find_libreoffice()
    if not soffice_path:
        raise LibreOfficeNotFoundError(
            "LibreOffice not found. Install it:\n"
            "  macOS: brew install --cask libreoffice\n"
            "  Ubuntu: sudo apt install libreoffice"
        )
    
    # Set output directory
    if output_dir is None:
        output_dir = input_path.parent
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
    
    try:
        # Use a temporary directory for LibreOffice user profile to avoid conflicts
        with tempfile.TemporaryDirectory() as temp_profile:
            cmd = [
                soffice_path,
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(output_dir),
                f"-env:UserInstallation=file://{temp_profile}",
                str(input_path),
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=timeout,
            )
            
            if result.returncode != 0:
                return {
                    "success": False,
                    "error": f"LibreOffice conversion failed: {result.stderr}",
                    "output_path": None,
                }
            
            # Determine output filename
            output_path = output_dir / f"{input_path.stem}.docx"
            
            if not output_path.exists():
                return {
                    "success": False,
                    "error": f"Conversion completed but output file not found at {output_path}",
                    "output_path": None,
                }
            
            return {
                "success": True,
                "output_path": str(output_path),
            }
            
    except subprocess.TimeoutExpired:
        return {
            "success": False,
            "error": f"Conversion timed out after {timeout} seconds",
            "output_path": None,
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "output_path": None,
        }


def convert_docx_to_pdf(
    input_path: Union[str, Path],
    output_dir: Optional[Union[str, Path]] = None,
    timeout: int = 120,
) -> Dict[str, Any]:
    """
    Convert a .docx file to PDF using LibreOffice.
    
    This provides high-fidelity PDF output that preserves the original
    Word document layout, fonts, and styling.
    
    Args:
        input_path: Path to the .docx file
        output_dir: Directory for output file (defaults to same as input)
        timeout: Maximum seconds to wait for conversion
        
    Returns:
        Dictionary with operation result:
        {
            "success": bool,
            "output_path": str,
            "error": str (if failed)
        }
        
    Raises:
        LibreOfficeNotFoundError: If LibreOffice is not installed
        ConversionError: If conversion fails
    """
    input_path = Path(input_path)
    
    if not input_path.exists():
        return {
            "success": False,
            "error": f"Input file not found: {input_path}",
            "output_path": None,
        }
    
    # Find LibreOffice
    soffice_path = _find_libreoffice()
    if not soffice_path:
        raise LibreOfficeNotFoundError(
            "LibreOffice not found. Install it:\n"
            "  macOS: brew install --cask libreoffice\n"
            "  Ubuntu: sudo apt install libreoffice"
        )
    
    # Set output directory
    if output_dir is None:
        output_dir = input_path.parent
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
    
    try:
        # Use a temporary directory for LibreOffice user profile to avoid conflicts
        with tempfile.TemporaryDirectory() as temp_profile:
            cmd = [
                soffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(output_dir),
                f"-env:UserInstallation=file://{temp_profile}",
                str(input_path),
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=timeout,
            )
            
            if result.returncode != 0:
                return {
                    "success": False,
                    "error": f"LibreOffice PDF conversion failed: {result.stderr}",
                    "output_path": None,
                }
            
            # Determine output filename
            output_path = output_dir / f"{input_path.stem}.pdf"
            
            if not output_path.exists():
                return {
                    "success": False,
                    "error": f"Conversion completed but PDF not found at {output_path}",
                    "output_path": None,
                }
            
            return {
                "success": True,
                "output_path": str(output_path),
            }
            
    except subprocess.TimeoutExpired:
        return {
            "success": False,
            "error": f"PDF conversion timed out after {timeout} seconds",
            "output_path": None,
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "output_path": None,
        }


# =============================================================================
# MAIN PIPELINE FUNCTION
# =============================================================================

def fill_and_export_template(
    template_path: Union[str, Path],
    output_path: Union[str, Path],
    context: Dict[str, Any],
    export_pdf: bool = True,
    pdf_output_dir: Optional[Union[str, Path]] = None,
) -> Dict[str, Any]:
    """
    Fill a Word template and optionally export to PDF.
    
    This is the main entry point for template-based document generation.
    It combines template rendering and PDF export into a single operation.
    
    Args:
        template_path: Path to the .docx template file
        output_path: Path for the filled .docx output
        context: Dictionary of variables to fill in the template
        export_pdf: If True, also export to PDF
        pdf_output_dir: Directory for PDF (defaults to same as docx)
        
    Returns:
        Dictionary with operation result:
        {
            "success": bool,
            "docx_path": str,
            "pdf_path": str or None,
            "placeholders_found": list,
            "placeholders_filled": list,
            "placeholders_unfilled": list,
            "error": str (if failed)
        }
        
    Example:
        result = fill_and_export_template(
            template_path="/templates/letter_of_rep.docx",
            output_path="/output/Wilson_LOR.docx",
            context={
                "client_name": "John Wilson",
                "date_of_accident": "January 15, 2024",
                "insurance_company": "State Farm",
                "claim_number": "CLM-123456"
            },
            export_pdf=True
        )
    """
    template_path = Path(template_path)
    output_path = Path(output_path)
    
    # Handle .doc templates - convert to .docx first
    if template_path.suffix.lower() == ".doc":
        with tempfile.TemporaryDirectory() as temp_dir:
            convert_result = convert_doc_to_docx(template_path, temp_dir)
            if not convert_result["success"]:
                return {
                    "success": False,
                    "error": f"Failed to convert .doc template: {convert_result.get('error')}",
                    "docx_path": None,
                    "pdf_path": None,
                }
            template_path = Path(convert_result["output_path"])
    
    # Render the template
    render_result = render_docx_template(
        template_path=template_path,
        context=context,
        output_path=output_path,
    )
    
    if not render_result.get("success"):
        return {
            "success": False,
            "error": render_result.get("error", "Template rendering failed"),
            "docx_path": None,
            "pdf_path": None,
        }
    
    result = {
        "success": True,
        "docx_path": str(output_path),
        "pdf_path": None,
        "placeholders_found": render_result.get("placeholders_found", []),
        "placeholders_filled": render_result.get("placeholders_filled", []),
        "placeholders_unfilled": render_result.get("placeholders_unfilled", []),
    }
    
    # Export to PDF if requested
    if export_pdf:
        pdf_result = convert_docx_to_pdf(
            input_path=output_path,
            output_dir=pdf_output_dir,
        )
        
        if pdf_result.get("success"):
            result["pdf_path"] = pdf_result["output_path"]
        else:
            # PDF conversion failed but docx was created
            result["pdf_error"] = pdf_result.get("error", "PDF conversion failed")
    
    return result


# =============================================================================
# COMMAND-LINE INTERFACE (for testing)
# =============================================================================

if __name__ == "__main__":
    import sys
    import json
    
    print("Word Template Pipeline - Dependency Check")
    print("=" * 50)
    
    deps = check_dependencies()
    print(f"docxtpl available: {deps['docxtpl_available']}")
    print(f"python-docx available: {deps['python_docx_available']}")
    print(f"LibreOffice available: {deps['libreoffice_available']}")
    if deps['libreoffice_path']:
        print(f"LibreOffice path: {deps['libreoffice_path']}")
    
    if deps['missing']:
        print(f"\nMissing dependencies: {', '.join(deps['missing'])}")
        sys.exit(1)
    else:
        print("\n✅ All dependencies available!")
        sys.exit(0)
